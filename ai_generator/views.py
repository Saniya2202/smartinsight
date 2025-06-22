from docx import Document
from django.http import HttpResponse

def download_cover_letter(request):
    # Example cover letter (replace with generated one from session or database)
    cover_letter = request.session.get('cover_letter_text', 'No cover letter found.')

    document = Document()
    document.add_heading('Cover Letter', 0)
    document.add_paragraph(cover_letter)

    # Create HTTP response with DOCX file
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=cover_letter.docx'
    document.save(response)
    return response
